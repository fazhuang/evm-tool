import io
import os
import datetime
from docx import Document
from fpdf import FPDF

# Project root path logic inside class to find font
FONT_NAME = "SimHei"

def generate_docx_stream(report: dict, filename: str) -> io.BytesIO:
    doc = Document()
    doc.add_heading(f"【AI特检】{filename} - 审查诊断报告", 0)
    
    risks = report.get("合规性风险", [])
    if risks:
        doc.add_heading("🔴 合规风险 (必须整改)", level=1)
        for i, r in enumerate(risks, 1):
            doc.add_heading(f"风险 {i}: {r.get('描述', '未知')}", level=2)
            p = doc.add_paragraph()
            p.add_run("⚡ AI 修复建议：").bold = True
            p.add_run(r.get('建议', '无'))
            
    logics = report.get("逻辑错误", [])
    if logics:
        doc.add_heading("🟡 逻辑异常 (自相矛盾)", level=1)
        for i, l in enumerate(logics, 1):
            doc.add_heading(f"异常 {i}: {l.get('描述', '未知')}", level=2)
            p = doc.add_paragraph()
            p.add_run("⚡ AI 修复建议：").bold = True
            p.add_run(l.get('建议', '无'))
            
    infos = report.get("核心信息", [])
    if infos:
        doc.add_heading("🔵 核心提取信息", level=1)
        for i, info in enumerate(infos, 1):
            doc.add_paragraph(f"{info.get('描述', '未知')}：{info.get('建议', '')}", style='List Bullet')
            
    p = doc.add_paragraph("\n")
    run = p.add_run("华招国际内部参阅专用")
    run.bold = True
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def generate_pdf_stream(project_name: str, manager: str) -> io.BytesIO:
    pdf = FPDF()
    pdf.add_page()
    
    # Priority load bundled open-source Chinese fonts
    # Root dir of backend is where SimHei.ttf is copied
    backend_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    font_path = os.path.join(backend_root, f"{FONT_NAME}.ttf")
    
    if os.path.exists(font_path):
        pdf.add_font(FONT_NAME, "", font_path)
        pdf.set_font(FONT_NAME, "", 14)
    else:
        pdf.set_font("Helvetica", "", 14)
        
    pdf.cell(200, 10, text=f"催办函 - {project_name}", align="C")
    pdf.ln(15)
    
    body_text = (
        f"尊敬的 {manager} 负责人：\n\n"
        f"您好！\n\n"
        f"系统检测到【{project_name}】项目的电子化进度低于 50%，"
        f"且距离开标日期已不足 3 天。为确保项目顺利推进，请您立即跟进处理，"
        f"加快电子化流程，并在系统内更新最新进度。\n\n"
        f"特此提醒！\n\n"
        f"自动预警系统\n"
        f"生成时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    pdf.multi_cell(0, 10, text=body_text)
    
    bio = io.BytesIO(pdf.output())
    bio.seek(0)
    return bio
