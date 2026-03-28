import io
import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def assemble_gansu_bidding_doc(form_data: dict) -> io.BytesIO:
    """Assembles a standardized Gansu Provincial bidding document based on form data."""
    doc = Document()
    
    # Standard Gansu Header Styling
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "甘肃省政府采购标准化招标文件 (2024 V1.0)"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Title Page
    doc.add_heading("招标文件", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n" * 2)
    
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    
    fields = [
        ("项目名称", form_data.get('project_name', '')),
        ("项目编号", form_data.get('project_code', '')),
        ("项目类型", form_data.get('project_category', '政府采购-服务')),
        ("预算金额", f"¥ {form_data.get('budget_amount', 0):,.2f}"),
        ("采购中心/代理机构", form_data.get('agent_name', '')),
        ("编制日期", datetime.datetime.now().strftime('%Y年%m月%d日'))
    ]
    
    for i, (label, value) in enumerate(fields):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)

    doc.add_page_break()
    
    # Chapter 1: 招标公告
    doc.add_heading("第一章 招标公告", level=1)
    p = doc.add_paragraph()
    p.add_run(f"受 {form_data.get('owner_name', '招标人')} 委托，{form_data.get('agent_name', '代理公司')} 对 {form_data.get('project_name', '项目')} 进行公开招标，欢迎符合资格条件的供应商参加。").bold = False
    
    # Chapter 2: 投标人资格要求 (Dynamic Mapping)
    doc.add_heading("第二章 投标人资格要求", level=1)
    quals = form_data.get('supplier_qualifications', [])
    if not quals:
        quals = ["满足《中华人民共和国政府采购法》第二十二条规定。", "落实政府采购政策需满足的资格要求：详见招标文件。"]
        
    for q in quals:
        doc.add_paragraph(q, style='List Bullet')
        
    # Chapter 3: 采购需求与技术参数 (The Core AI-Gen Content)
    doc.add_heading("第三章 采购需求", level=1)
    doc.add_paragraph("根据标书制作人员提供的技术参数，AI 已自动生成如下采购需求。请代理机构核对：")
    
    specs = form_data.get('technical_specs', [])
    if specs:
        spec_table = doc.add_table(rows=1, cols=3)
        spec_table.style = 'Table Grid'
        hdr_cells = spec_table.rows[0].cells
        hdr_cells[0].text = '序号'
        hdr_cells[1].text = '参数项'
        hdr_cells[2].text = '具体要求'
        
        for i, spec in enumerate(specs, 1):
            row_cells = spec_table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = spec.get('name', '指标')
            row_cells[2].text = spec.get('value', '见附件')
    else:
        doc.add_paragraph("（技术参数详见甲方提供的正式参数清单附件）")

    # Chapter 4: 评标方法
    doc.add_heading("第四章 评标办法", level=1)
    method = form_data.get('evaluation_method', '综合评分法')
    doc.add_paragraph(f"本项目采用【{method}】。评分因素包括价格、技术、商务及售后服务。")
    
    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "本文件由 HZ-SAGE 招标文件智能制作系统生成"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
